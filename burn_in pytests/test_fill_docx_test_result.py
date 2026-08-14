# test_fill_docx_test_result.py
from pathlib import Path
import pytest
from docx import Document
from burn_in_module import fill_docx_test_result


@pytest.fixture
def docx_factory(tmp_path: Path):
    """Factory fixture to create temporary DOCX files for testing."""
    def _create(paragraphs: list[str | list[str]], filename: str = "template.docx") -> Path:
        file_path = tmp_path / filename
        doc = Document()
        for p_data in paragraphs:
            p = doc.add_paragraph()
            if isinstance(p_data, list):
                for run_text in p_data:
                    p.add_run(run_text)
            else:
                p.add_run(p_data)
        doc.save(file_path)
        return file_path
    return _create


@pytest.mark.parametrize(
    "paragraphs, test_result, expected_filled, expected_texts",
    [
        # --- Standard Replacement Behavior ---
        (
            ["Test Result: _____"],
            "pass",
            True,
            ["Test Result: PASS"],
        ),
        (
            ["Header Info", "Test Result: _______", "Footer Info"],
            "fail",
            True,
            ["Header Info", "Test Result: FAIL", "Footer Info"],
        ),
        (
            ["Test Result: " + "_" * 20],
            "conditional pass",
            True,
            ["Test Result: CONDITIONAL PASS"],
        ),
        (
            ["Test Result: _"],
            "n/a",
            True,
            ["Test Result: N/A"],
        ),
        # --- Run Fragmentation & Formatting Variants ---
        (
            [["Test Result: ", "_____"]],
            "pass",
            True,
            ["Test Result: PASS"],
        ),
        (
            [["Test Result:", " ", "___"]],
            "fail",
            True,
            ["Test Result: FAIL"],
        ),
        # --- Unmatched / No-Op Scenarios ---
        (
            ["Test Result: PASSED"],
            "pass",
            False,
            ["Test Result: PASSED"],
        ),
        (
            ["Customer Name: _________", "Date: _________"],
            "pass",
            False,
            ["Customer Name: _________", "Date: _________"],
        ),
        (
            ["No matching placeholder here"],
            "fail",
            False,
            ["No matching placeholder here"],
        ),
    ],
    ids=[
        "simple-pass",
        "multiline-with-context",
        "multi-word-result",
        "single-char-result",
        "split-across-runs",
        "split-label-and-underscore",
        "already-populated",
        "placeholders-without-target-label",
        "no-placeholders",
    ],
)
def test_fill_docx_test_result_behavior(
    docx_factory,
    tmp_path: Path,
    paragraphs: list[str | list[str]],
    test_result: str,
    expected_filled: bool,
    expected_texts: list[str],
):
    """
    Verifies that fill_docx_test_result replaces placeholders with uppercase results
    and returns correct status flags without corrupting surrounding document structure.
    """
    input_path = docx_factory(paragraphs)
    output_path = tmp_path / "output.docx"

    filled = fill_docx_test_result(str(input_path), test_result, str(output_path))

    assert filled is expected_filled
    assert output_path.exists()

    result_doc = Document(output_path)
    actual_texts = [p.text for p in result_doc.paragraphs]
    assert actual_texts == expected_texts


@pytest.mark.parametrize("in_place", [True, False], ids=["in-place-overwrite", "separate-output-file"])
def test_fill_docx_test_result_file_handling(docx_factory, tmp_path: Path, in_place: bool):
    """Verifies that the document can be written to a new path or overwritten in-place."""
    input_path = docx_factory(["Test Result: _____"])
    output_path = input_path if in_place else tmp_path / "destination.docx"

    filled = fill_docx_test_result(str(input_path), "pass", str(output_path))

    assert filled is True
    assert output_path.exists()
    assert Document(output_path).paragraphs[0].text == "Test Result: PASS"


def test_fill_docx_test_result_raises_error_for_missing_file(tmp_path: Path):
    """Ensures a FileNotFoundError or Document reading exception is propagated when given an invalid path."""
    non_existent_file = tmp_path / "missing.docx"
    output_file = tmp_path / "out.docx"

    with pytest.raises(Exception):
        fill_docx_test_result(str(non_existent_file), "PASS", str(output_file))